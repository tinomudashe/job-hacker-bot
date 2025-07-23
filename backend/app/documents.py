from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException, status
from sqlalchemy import desc
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.db import get_db
from app.models_db import Document, User, Resume
from app.dependencies import get_current_active_user
from app.cv_processor import cv_processor, CVExtractionResult
from app.enhanced_memory import EnhancedMemoryManager
from datetime import datetime
import os
import uuid
import asyncio
from pathlib import Path
import io
import docx
from pypdf import PdfReader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LCDocument
from typing import List, Optional
from pydantic import BaseModel
import shutil
import logging

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)



async def _update_user_profile_from_cv(user: User, cv_data: CVExtractionResult, db: AsyncSession) -> tuple[bool, list[str]]:
    """Updates user profile and resume data from extracted CV content if fields are not already set."""
    updated_fields = []

    # --- 1. Update User Table Fields ---
    field_map = {
        'full_name': 'name',
        'email': 'email',
        'phone_number': 'phone',
        'location': 'address',
        'linkedin_url': 'linkedin',
        'summary': 'profile_headline',
        'skills': 'skills'
    }

    for cv_key, user_key in field_map.items():
        cv_value = getattr(cv_data, cv_key, None)
        if cv_value:
            if isinstance(cv_value, list):
                cv_value = ", ".join(skill for skill in cv_value if skill)
            if not getattr(user, user_key, None):
                setattr(user, user_key, cv_value)
                updated_fields.append(user_key)
    
    if cv_data.full_name and (not user.first_name or not user.last_name):
        parts = cv_data.full_name.split()
        if len(parts) > 1:
            if not user.first_name:
                user.first_name = parts[0]
                updated_fields.append('first_name')
            if not user.last_name:
                user.last_name = " ".join(parts[1:])
                updated_fields.append('last_name')
        elif not user.first_name:
            user.first_name = cv_data.full_name
            updated_fields.append('first_name')

    # --- 2. Update Resume JSON Data (Experience & Education) ---
    resume_result = await db.execute(select(Resume).where(Resume.user_id == user.id))
    db_resume = resume_result.scalars().first()

    if not db_resume:
        db_resume = Resume(user_id=user.id, data={})
        db.add(db_resume)
        updated_fields.append('resume_record_created')
    
    resume_data = db_resume.data if db_resume.data else {}

    # Update Work Experience if it's empty
    if cv_data.experience and not resume_data.get('experience'):
        resume_data['experience'] = [exp.dict() for exp in cv_data.experience]
        updated_fields.append('work_experience')

    # Update Education if it's empty
    if cv_data.education and not resume_data.get('education'):
        resume_data['education'] = [edu.dict() for edu in cv_data.education]
        updated_fields.append('education')

    # Update resume data field
    if 'work_experience' in updated_fields or 'education' in updated_fields or 'resume_record_created' in updated_fields:
        db_resume.data = resume_data

    if updated_fields:
        db.add(user)
        return True, updated_fields
    return False, []

router = APIRouter()
logger = logging.getLogger(__name__)

class DocumentOut(BaseModel):
    id: str
    type: str
    name: str
    date_created: datetime
    date_updated: datetime

    class Config:
        from_attributes = True

class CVUploadResponse(BaseModel):
    document: DocumentOut
    extracted_info: Optional[CVExtractionResult] = None
    profile_updated: bool = False
    auto_extracted_fields: List[str] = []
    personalized_insights: Optional[str] = None


@router.get("/documents/resumes/latest", response_model=dict, summary="Get the latest generated resume")
async def get_latest_resume(
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Retrieves the most recently generated resume for the authenticated user.
    It checks the 'generated_content' table for entries that are resumes.
    """
    logger.info(f"Fetching latest generated resume for user {db_user.id}")
    
    # NOTE: Your generated resumes are currently saved in the 'GeneratedCoverLetter' table.
    # This query targets that table. We will identify resumes by looking for a marker.
    from app.models_db import GeneratedCoverLetter
    
    result = await db.execute(
        select(GeneratedCoverLetter)
        .where(GeneratedCoverLetter.user_id == db_user.id)
        # A simple way to distinguish resumes from cover letters in the same table
        .where(GeneratedCoverLetter.content.contains("[DOWNLOADABLE_RESUME]"))
        .order_by(desc(GeneratedCoverLetter.created_at))
        .limit(1)
    )
    
    latest_resume = result.scalars().first()
    
    if not latest_resume:
        raise HTTPException(status_code=404, detail="No generated resume found.")
        
    # The content is a string containing markers and the resume text.
    # We will return the whole object so the frontend can parse it.
    return {
        "id": latest_resume.id,
        "user_id": latest_resume.user_id,
        "content": latest_resume.content,
        "created_at": latest_resume.created_at
    }




@router.post("/documents", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Form(...),  # 'resume', 'cover_letter', or 'generic'
    name: str = Form(...),
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    if doc_type not in ("resume", "cover_letter", "generic"):
        raise HTTPException(status_code=400, detail="Invalid document type. Must be 'resume', 'cover_letter', or 'generic'.")
    
    logger.info(f"Starting document upload: {name} ({doc_type}) for user {db_user.id}")
    
    # Save file
    doc_id = str(uuid.uuid4())
    user_dir = UPLOAD_DIR / db_user.id
    user_dir.mkdir(exist_ok=True)
    file_path = user_dir / f"{doc_id}_{file.filename}"
    
    try:
        with open(file_path, "wb") as f:
            f.write(await file.read())
        
        # Extract text
        if file.filename.lower().endswith(".pdf"):
            try:
                pdf = PdfReader(str(file_path))
                text = "".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error processing PDF: {e}")
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading file: {e}")
        
        # Build FAISS index with enhanced embeddings
        embedding = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
        
        # Smart chunking based on document type and user preferences
        chunk_size = _get_optimal_chunk_size(doc_type, len(text))
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        docs = [LCDocument(page_content=chunk) for chunk in chunks]
        
        faiss_dir = user_dir / f"faiss_{doc_id}"
        faiss_dir.mkdir(exist_ok=True)
        
        # Use async FAISS methods to prevent greenlet errors
        vectorstore = await FAISS.afrom_documents(docs, embedding)
        
        # Run save_local in thread executor to avoid blocking
        await asyncio.get_event_loop().run_in_executor(
            None, vectorstore.save_local, str(faiss_dir)
        )
        
        # Store document metadata FIRST (avoid transaction conflicts)
        now = datetime.utcnow()
        doc = Document(
            id=doc_id,
            user_id=db_user.id,
            type=doc_type,
            name=name,
            content=text,
            vector_store_path=str(faiss_dir),
            date_created=now,
            date_updated=now
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        
        return doc
        
    except Exception as e:
        # Ensure transaction rollback on main error
        try:
            if db.is_active:
                await db.rollback()
        except Exception:
            pass
            
        logger.error(f"Error processing document upload: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")
    

    
@router.post("/upload", summary="Upload a generic file", status_code=status.HTTP_201_CREATED)
async def upload_file(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user),
):
    """
    Handles generic file uploads. Saves the file to the server and creates a
    database entry for it.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided.")

    file_id = str(uuid.uuid4())
    file_extension = Path(file.filename).suffix
    new_filename = f"{file_id}{file_extension}"
    file_path = UPLOAD_DIR / new_filename

    try:
        # Save file to disk
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())

        # Create a new document record in the database
        new_document = Document(
            id=file_id,
            user_id=db_user.id,
            name=file.filename,
            path=str(file_path),
            file_type=file.content_type,
            size=file_path.stat().st_size,
            created_at=datetime.utcnow(),
        )
        db.add(new_document)
        await db.commit()
        await db.refresh(new_document)

        logger.info(f"User {db_user.id} uploaded file: {new_document.name}")

        return {
            "filename": new_document.name,
            "document_id": new_document.id,
            "size": new_document.size,
            "file_type": new_document.file_type,
        }
    except Exception as e:
        # Clean up the file if the database operation fails
        if file_path.exists():
            os.remove(file_path)
        logger.error(f"Error during file upload for user {db_user.id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred during file upload: {e}",
        )


@router.post("/documents/cv-upload", response_model=CVUploadResponse, status_code=status.HTTP_201_CREATED)
async def upload_cv_with_extraction(
    file: UploadFile = File(...),
    name: str = Form(...),
    auto_update_profile: bool = Form(True),
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Simplified CV upload without enhanced memory features to avoid greenlet errors
    """
    
    try:
        # Save file
        doc_id = str(uuid.uuid4())
        user_dir = UPLOAD_DIR / db_user.id
        user_dir.mkdir(exist_ok=True)
        file_path = user_dir / f"{doc_id}_{file.filename}"
        
        with open(file_path, "wb") as f:
            f.write(await file.read())
        
        # Enhanced text extraction
        if file.filename and file.filename.lower().endswith(".pdf"):
            try:
                pdf = PdfReader(str(file_path))
                text = "".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error processing PDF: {e}")
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading file: {e}")
        
        # Store document metadata with basic FAISS processing
        now = datetime.utcnow()
        doc = Document(
            id=doc_id,
            user_id=db_user.id,
            type="resume",
            name=name,
            content=text,
            vector_store_path=None,  # Will be added after FAISS processing
            date_created=now,
            date_updated=now
        )
        db.add(doc)
        await db.commit()
        # Skip refresh to avoid greenlet issues - we have all the data we need
        
        # Now, process CV data and update profile in a separate step
        # This separation helps isolate transaction scopes and debug greenlet issues
        try:
            # FIX: Call the correct method 'extract_cv_information' and pass the file_path.
            cv_data = await cv_processor.extract_cv_information(file_path)
            
            # Auto-update profile if enabled
            profile_updated = False
            updated_fields = []
            if auto_update_profile:
                profile_updated, updated_fields = await _update_user_profile_from_cv(db_user, cv_data, db)
                if profile_updated:
                    await db.commit()
                    await db.refresh(db_user) # Refresh user to get latest state
            
            # This is a placeholder for the more advanced features you were working on
            insights = "This is a placeholder for personalized insights."
            
            # Construct final response
            return CVUploadResponse(
                document=doc,
                extracted_info=cv_data,
                profile_updated=profile_updated,
                auto_extracted_fields=updated_fields,
                personalized_insights=insights
            )
            
        except Exception as e:
            # Handle errors during CV processing specifically
            logger.error(f"Error during CV processing stage for doc {doc.id}: {e}")
            # The document is already saved, so we return the doc info with an error message
            # This is better than rolling back the whole upload
            return CVUploadResponse(
                document=doc,
                extracted_info=None,
                profile_updated=False,
                auto_extracted_fields=[],
                personalized_insights=f"CV processing failed: {e}"
            )

    except Exception as e:
        logger.error(f"Major error during CV upload for user {db_user.id}: {e}")
        raise HTTPException(status_code=500, detail=f"A critical error occurred: {e}")


async def _analyze_document_with_context(
    content: str, 
    doc_type: str, 
    user_profile, 
    memory_manager: Optional[EnhancedMemoryManager]
) -> dict:
    # This is a placeholder for the advanced analysis function you were building
    # It would use the user profile and memory for deeper insights
    if doc_type == "resume":
        return await _analyze_resume_content(content, user_profile)
    elif doc_type == "cover_letter":
        return await _analyze_cover_letter_content(content, user_profile)
    return {}

async def _generate_cv_insights_with_context(
    content: str,
    user_profile,
    memory_manager: Optional[EnhancedMemoryManager]
) -> Optional[str]:
    # Placeholder for advanced insight generation
    
    # Example of how you might use the context:
    if memory_manager and user_profile:
        past_feedback = await memory_manager.get_related_memories(
            "cv_feedback",
            limit=3,
            min_similarity=0.78
        )
        if past_feedback:
            # Generate insights based on past interactions
            # This is a simplified example
            return "Based on your past CVs, you've improved your project descriptions. Focus next on quantifying achievements."

    # Fallback to basic analysis if no context is available
    themes = _extract_resume_themes(content)
    if not themes:
        return None
    return f"This CV focuses on themes of: {', '.join(themes)}. Consider tailoring this to the job description."

async def _learn_from_cv_content(content: str, memory_manager: Optional[EnhancedMemoryManager]):
    if not memory_manager:
        return

    # Extract key entities and learn them
    skills = _extract_skills_from_cv(content)
    if skills:
        await memory_manager.add_memory(
            "user_skill", 
            metadata={"skills": skills, "source": "cv_upload"}
        )

    experience_years = _estimate_experience_years(content)
    if experience_years is not None:
        await memory_manager.add_memory(
            "user_experience_level", 
            metadata={"years": experience_years, "source": "cv_upload"}
        )

    industry = _extract_industry_keywords(content)
    if industry:
        await memory_manager.add_memory(
            "user_industry_preference",
            metadata={"industry": industry, "source": "cv_upload"}
        )

    logger.info("Learned skills, experience, and industry from CV content.")


def _get_optimal_chunk_size(doc_type: str, content_length: int) -> int:
    # Simple logic to adjust chunk size
    if doc_type == "resume":
        return 512  # Smaller chunks for denser resume content
    elif content_length > 10000:
        return 2048 # Larger chunks for very long documents
    return 1024

def _extract_resume_themes(content: str) -> List[str]:
    # This is a simplified keyword-based theme extraction
    themes = []
    content_lower = content.lower()
    if "lead" in content_lower or "manage" in content_lower:
        themes.append("Leadership")
    if "python" in content_lower or "java" in content_lower or "react" in content_lower:
        themes.append("Software Development")
    if "data" in content_lower or "analyst" in content_lower:
        themes.append("Data Analysis")
    if "ui/ux" in content_lower or "design" in content_lower:
        themes.append("User Experience")
    return list(set(themes))

def _extract_cover_letter_themes(content: str) -> List[str]:
    themes = []
    content_lower = content.lower()
    if "excited" in content_lower or "passionate" in content_lower:
        themes.append("Enthusiasm")
    if "perfect fit" in content_lower or "well-suited" in content_lower:
        themes.append("Fit for Role")
    if "look forward" in content_lower:
        themes.append("Proactive Follow-up")
    return list(set(themes))

def _extract_skills_from_cv(content: str) -> List[str]:
    # A more robust regex could be used here
    # This is a simple example
    skills = [
        "Python", "Java", "React", "TypeScript", "JavaScript", "SQL",
        "AWS", "GCP", "Docker", "Kubernetes", "Spring Boot"
    ]
    found_skills = [skill for skill in skills if skill.lower() in content.lower()]
    return list(set(found_skills))

def _estimate_experience_years(content: str) -> Optional[int]:
    import re
    # Look for patterns like "X+ years", "X years of experience"
    matches = re.findall(r'(\d+)\+?\s*years', content, re.IGNORECASE)
    if matches:
        return max(int(m) for m in matches)
    
    # A more complex method would parse start/end dates of jobs
    # This is a simplified version for demonstration
    
    # Count occurrences of year numbers to make a rough guess
    year_mentions = re.findall(r'\b(20\d{2})\b', content)
    if len(year_mentions) > 1:
        try:
            min_year = min(int(y) for y in year_mentions)
            max_year = max(int(y) for y in year_mentions)
            # Avoid unrealistic ranges
            if 0 < (max_year - min_year) < 40:
                return max_year - min_year
        except ValueError:
            pass

    return None

def _extract_industry_keywords(content: str) -> Optional[str]:
    content_lower = content.lower()
    if "finance" in content_lower or "fintech" in content_lower:
        return "Finance"
    if "healthcare" in content_lower or "medical" in content_lower:
        return "Healthcare"
    if "ecommerce" in content_lower or "retail" in content_lower:
        return "E-commerce"
    if "saas" in content_lower:
        return "SaaS"
    return None


@router.get("/documents", response_model=List[DocumentOut])
async def list_documents(
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """Get all documents for the authenticated user."""
    logger.info(f"Listing documents for user {db_user.id}")
    result = await db.execute(select(Document).where(Document.user_id == db_user.id))
    return result.scalars().all()

@router.delete("/documents/delete-all", status_code=status.HTTP_204_NO_CONTENT)
async def delete_all_documents_for_user(
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Deletes all documents (and their associated FAISS indexes) for the authenticated user.
    This is a destructive operation.
    """
    logger.warning(f"Initiating deletion of all documents for user {db_user.id}")
    
    # First, get all documents to find their FAISS index paths
    docs_result = await db.execute(select(Document).where(Document.user_id == db_user.id))
    documents_to_delete = docs_result.scalars().all()
    
    deleted_count = 0
    for doc in documents_to_delete:
        # Delete the physical FAISS index directory if it exists
        if doc.vector_store_path and os.path.exists(doc.vector_store_path):
            try:
                shutil.rmtree(doc.vector_store_path)
                logger.info(f"Deleted FAISS index for document {doc.id} at {doc.vector_store_path}")
            except Exception as e:
                logger.error(f"Error deleting FAISS index for document {doc.id}: {e}")
        
        # Delete the document record from the database
        await db.delete(doc)
        deleted_count += 1
        
    if deleted_count > 0:
        try:
            await db.commit()
            logger.info(f"Successfully deleted {deleted_count} documents for user {db_user.id}")
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to commit deletion of documents for user {db_user.id}: {e}")
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Database error during document deletion.")
    else:
        # If no documents were found, we can just return success
        logger.info(f"No documents found to delete for user {db_user.id}")

@router.get("/documents/{doc_id}", response_model=DocumentOut)
async def get_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """Get a single document by its ID."""
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == db_user.id))
    doc = result.scalars().first()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return doc

@router.post("/documents/{doc_id}/reprocess-cv")
async def reprocess_cv_extraction(
    doc_id: str,
    auto_update_profile: bool = True,
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Manually re-trigger CV processing for an existing document.
    This is useful if the initial extraction failed or needs to be updated.
    """
    logger.info(f"Reprocessing CV for doc {doc_id}, user {db_user.id}")
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == db_user.id))
    doc = result.scalars().first()
    if not doc or doc.type != 'resume':
        raise HTTPException(status_code=404, detail="Resume document not found")

    if not doc.content:
        raise HTTPException(status_code=400, detail="Document has no content to process")

    try:
        # Reconstruct the file path. This assumes the file exists at the path stored in the DB.
        # A more robust solution would check if the file exists.
        file_path = Path(doc.path) if doc.path else None
        if not file_path or not file_path.exists():
            # If path is not stored or file deleted, create a temporary file to process
            user_dir = UPLOAD_DIR / db_user.id
            user_dir.mkdir(exist_ok=True)
            # Use doc.id to avoid filename collisions
            temp_file_path = user_dir / f"temp_{doc.id}.tmp"
            with open(temp_file_path, "w", encoding="utf-8") as f:
                f.write(doc.content)
            file_path = temp_file_path
        
        # FIX: Call the correct method 'extract_cv_information'.
        cv_data = await cv_processor.extract_cv_information(file_path)
        
        # Clean up temporary file if it was created
        if 'temp_file_path' in locals() and temp_file_path.exists():
            os.remove(temp_file_path)

        profile_updated = False
        updated_fields = []
        if auto_update_profile:
            profile_updated, updated_fields = await _update_user_profile_from_cv(db_user, cv_data, db)
            if profile_updated:
                await db.commit()
                await db.refresh(db_user)

        return {
            "message": "CV re-processed successfully.",
            "extracted_info": cv_data,
            "profile_updated": profile_updated,
            "auto_extracted_fields": updated_fields,
        }
    except Exception as e:
        logger.error(f"Error during manual CV reprocessing for doc {doc.id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to reprocess CV: {e}")

@router.delete("/documents/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """Delete a single document by its ID."""
    logger.info(f"Attempting to delete document {doc_id} for user {db_user.id}")
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == db_user.id))
    doc = result.scalars().first()
    
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    
    # Delete the physical FAISS index directory if it exists
    if doc.vector_store_path and os.path.exists(doc.vector_store_path):
        try:
            shutil.rmtree(doc.vector_store_path)
            logger.info(f"Deleted FAISS index at {doc.vector_store_path}")
        except Exception as e:
            # Log error but don't block deletion of the DB record
            logger.error(f"Error deleting FAISS index for doc {doc.id}: {e}")
    
    try:
        await db.delete(doc)
        await db.commit()
        logger.info(f"Successfully deleted document {doc_id} for user {db_user.id}")
    except Exception as e:
        await db.rollback()
        logger.error(f"Database error deleting document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete document from database.")


class DocumentUpdate(BaseModel):
    name: Optional[str] = None
    type: Optional[str] = None

@router.put("/documents/{doc_id}", response_model=DocumentOut)
async def update_document(
    doc_id: str,
    doc_update: DocumentUpdate,
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """Update a document's metadata (name or type)."""
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == db_user.id))
    doc = result.scalars().first()
    
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
        
    update_data = doc_update.dict(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="No update data provided.")

    for key, value in update_data.items():
        setattr(doc, key, value)
    
    doc.date_updated = datetime.utcnow()
    
    try:
        await db.commit()
        await db.refresh(doc)
        return doc
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update document: {e}")

@router.get("/documents/insights", response_model=dict)
async def get_personalized_document_insights(
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Generate comprehensive insights across all of the user's documents.
    This provides a holistic view of their professional profile.
    """
    logger.info(f"Generating comprehensive insights for user {db_user.id}")
    
    # Get user profile for context
    # This is a simplified user profile object for demonstration
    user_profile = {
        "name": db_user.name,
        "profile_headline": db_user.profile_headline,
        "skills": db_user.skills.split(',') if db_user.skills else []
    }

    docs_result = await db.execute(select(Document).where(Document.user_id == db_user.id))
    documents = docs_result.scalars().all()
    
    if not documents:
        return {"message": "No documents found to generate insights."}

    # Initialize memory manager (optional, for advanced context)
    memory_manager = None
    if db_user.faiss_index_path and os.path.exists(db_user.faiss_index_path):
        try:
            memory_manager = EnhancedMemoryManager(user_id=db_user.id, db_session=db)
            await memory_manager.load_memory()
        except Exception as e:
            logger.warning(f"Could not load enhanced memory for user {db_user.id}: {e}")

    try:
        insights = await _generate_comprehensive_document_insights(documents, user_profile, memory_manager)
        return insights
    except Exception as e:
        logger.error(f"Error generating insights for user {db_user.id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate document insights.")

@router.get("/documents/{doc_id}/analysis", response_model=dict)
async def get_document_analysis(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    db_user: User = Depends(get_current_active_user)
):
    """
    Perform a detailed analysis of a single document, providing specific feedback.
    """
    logger.info(f"Analyzing document {doc_id} for user {db_user.id}")
    
    # Get user profile for context
    user_profile = {
        "name": db_user.name,
        "profile_headline": db_user.profile_headline,
        "skills": db_user.skills.split(',') if db_user.skills else []
    }

    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == db_user.id))
    doc = result.scalars().first()
    
    if not doc:
        raise HTTPException(status_code=4.04, detail="Document not found")

    # Initialize memory manager for contextual analysis
    memory_manager = None
    if db_user.faiss_index_path and os.path.exists(db_user.faiss_index_path):
        try:
            memory_manager = EnhancedMemoryManager(user_id=db_user.id, db_session=db)
            await memory_manager.load_memory()
        except Exception as e:
            logger.warning(f"Could not load enhanced memory for single doc analysis: {e}")

    try:
        analysis = await _analyze_single_document(doc, user_profile, memory_manager)
        return analysis
    except Exception as e:
        logger.error(f"Error analyzing document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to analyze document.")


# --- Helper Functions for Analysis ---

async def _generate_comprehensive_document_insights(
    documents: List[Document],
    user_profile,
    memory_manager: Optional[EnhancedMemoryManager]
) -> dict:
    """
    Analyzes multiple documents to find overarching themes, skill clusters,
    and potential inconsistencies.
    """
    if not documents:
        return {}

    all_content = " ".join([doc.content for doc in documents if doc.content])
    
    # High-level summary
    summary = {
        "total_documents": len(documents),
        "document_types": {t: len([d for d in documents if d.type == t]) for t in set(d.type for d in documents)},
        "overall_themes": _extract_resume_themes(all_content) + _extract_cover_letter_themes(all_content),
        "identified_skills": _extract_skills_from_cv(all_content),
        "estimated_experience_years": _estimate_experience_years(all_content),
        "suggested_industries": _extract_industry_keywords(all_content)
    }

    # Example of using memory for deeper insights
    if memory_manager:
        related_jobs = await memory_manager.get_related_memories("job_application", limit=5)
        if related_jobs:
            summary["related_past_applications"] = [mem.metadata.get("job_title") for mem in related_jobs]

    # Suggestions for improvement
    suggestions = []
    if len(summary["identified_skills"]) < 10:
        suggestions.append("Consider adding more specific technical or soft skills to your documents.")
    if not summary["estimated_experience_years"]:
        suggestions.append("Your documents don't clearly state your years of experience. Consider adding a summary statement like 'Software Engineer with 5+ years of experience'.")
    if "Leadership" not in summary["overall_themes"]:
        suggestions.append("To target senior roles, try to incorporate language that highlights leadership or mentorship experience.")

    return {
        "summary": summary,
        "suggestions_for_improvement": suggestions
    }


async def _analyze_single_document(
    document: Document,
    user_profile,
    memory_manager: Optional[EnhancedMemoryManager]
) -> dict:
    """
    Performs a deep dive into one document.
    """
    if not document.content:
        return {"error": "Document has no content to analyze."}

    analysis = {
        "doc_id": document.id,
        "doc_name": document.name,
        "doc_type": document.type,
        "analysis_results": {}
    }

    if document.type == "resume":
        analysis["analysis_results"] = await _analyze_resume_content(document.content, user_profile)
    elif document.type == "cover_letter":
        analysis["analysis_results"] = await _analyze_cover_letter_content(document.content, user_profile)
    else:
        analysis["analysis_results"] = {
            "content_length": len(document.content),
            "preview": document.content[:200] + "..."
        }

    # Example of using memory for contextual feedback
    if memory_manager:
        job_context = await memory_manager.get_latest_memory("job_application_context")
        if job_context:
            # This is a placeholder for a more complex analysis
            # e.g., compare document skills to job description skills
            analysis["contextual_feedback"] = {
                "message": "Analysis based on your recent activity.",
                "related_job_title": job_context.metadata.get("job_title"),
                "suggestion": "This document seems well-aligned for a senior role. Consider emphasizing your project management skills more."
            }

    return analysis

async def _analyze_resume_content(content: str, user_profile) -> dict:
    """Specific analysis for resume documents."""
    
    # Basic metrics
    word_count = len(content.split())
    
    # Content analysis
    skills = _extract_skills_from_cv(content)
    themes = _extract_resume_themes(content)
    years_experience = _estimate_experience_years(content)

    # Action verb check (simplified)
    action_verbs = ["developed", "led", "managed", "created", "implemented", "designed"]
    action_verb_count = sum(1 for verb in action_verbs if verb in content.lower())

    # Suggestions for improvement
    optimization_tips = _generate_resume_optimization_tips(content, user_profile)

    return {
        "word_count": word_count,
        "detected_skills": skills,
        "primary_themes": themes,
        "estimated_years_experience": years_experience,
        "action_verb_score": f"{action_verb_count}/{len(action_verbs)}",
        "optimization_tips": optimization_tips
    }

async def _analyze_cover_letter_content(content: str, user_profile) -> dict:
    """Specific analysis for cover letter documents."""
    
    word_count = len(content.split())
    themes = _extract_cover_letter_themes(content)

    # Personalization check (simplified)
    # A real implementation would check for a company name
    is_personalized = "company" in content.lower() or "team" in content.lower()

    # Call to action check
    has_call_to_action = "look forward to hearing from you" in content.lower() or "discuss my qualifications" in content.lower()

    return {
        "word_count": word_count,
        "primary_themes": themes,
        "is_personalized": is_personalized,
        "has_clear_call_to_action": has_call_to_action,
        "suggestions": [
            "Ensure you mention the specific company and role you're applying for.",
            "Try to connect your skills directly to the requirements in the job description.",
            "End with a confident closing and a clear call to action."
        ]
    }

def _generate_resume_optimization_tips(content: str, user_profile) -> List[str]:
    """Generate ATS-friendly and other optimization tips."""
    tips = []
    
    # Tip 1: Quantify achievements
    if not any(char.isdigit() for char in content):
        tips.append("Quantify your achievements. Instead of 'improved performance', use 'improved performance by 15%'.")

    # Tip 2: Use action verbs
    if "responsible for" in content.lower():
        tips.append("Replace 'responsible for' with strong action verbs like 'managed', 'developed', or 'led'.")

    # Tip 3: Tailor skills
    user_skills = user_profile.get("skills", [])
    content_skills = _extract_skills_from_cv(content)
    if user_skills and not all(skill in content_skills for skill in user_skills[:3]):
        tips.append(f"Tailor your skills section for each job. Ensure top skills like {', '.join(user_skills[:3])} are present if relevant.")

    return tips

def _get_industry_keywords(industry: str) -> List[str]:
    # Example keyword list for different industries
    keywords = {
        "Finance": ["fintech", "trading", "risk", "compliance", "asset management"],
        "Healthcare": ["hipaa", "emr", "telehealth", "medical imaging", "patient data"],
        "E-commerce": ["checkout", "sku", "supply chain", "user conversion", "inventory"]
    }
    return keywords.get(industry, [])